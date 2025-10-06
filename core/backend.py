
# -*- coding: utf-8 -*-
"""
Backend (capa de datos y utilidades)
- Detección flexible de columnas (mapeo)
- Normalización y parseo de fechas (incluye números de serie Excel)
- Utilidades para manipular la tabla del DOCX (python-docx) con selección de tabla
- Lectura de placeholders por grupo desde Excel
"""
from __future__ import annotations

import io
import re
from datetime import datetime
from typing import Dict, Optional, List, Tuple

import pandas as pd
from docx import Document
from unidecode import unidecode

# ================= Normalización / utilidades ================= #

def _norm(s: str) -> str:
    s = unidecode(str(s or "").strip().lower())
    s = re.sub(r"\s+", " ", s)
    return s

def slugify(text: str) -> str:
    s = unidecode(str(text or "").strip())
    s = re.sub(r"[^A-Za-z0-9_\- ]+", "", s)
    s = s.strip().replace(" ", "_")
    return s or "SIN_GRUPO"

# ================= Columnas esperadas / sinónimos ================= #

SYNONYMS = {
    "actor": [
        "actor", "columna a", "a", "interesado", "responsable", "nombre del actor"
    ],
    "grupo": [
        "dependencia", "secretaria", "secretaría", "secretaria de", "entidad", "despacho",
        "direccion", "dirección", "institucion", "institución", "dependencia/entidad",
        "institucional", "grupo", "responsable"
    ],
    "mesa": [
        "nombre de la mesa", "nombre mesa", "mesa", "tema", "asunto", "nombre mesa/tema",
        "nombre", "actividad"
    ],
    "nivel": [
        "nivel", "nivel de la mesa", "compromiso", "tipo", "categoria", "categoría"
    ],
    "fecha": [
        "fecha", "fecha mesa", "fecha programada", "dia", "día", "dia mesa",
        "fecha de realizacion", "fecha de realización", "fecha programada mesa"
    ],
    "dato": [
        "dato transformador", "dato", "transformador", "descripcion dato",
        "descripción dato", "datos transformadores", "resultado esperado"
    ],
}

EXPECTED_HEADERS = [
    "nombre de la mesa", "nivel", "fecha", "dato transformador"
]

# ================= Mapeo de columnas ================= #

def guess_mapping(df: pd.DataFrame) -> Dict[str, Optional[str]]:
    """Retorna mapeo {'actor','mesa','nivel','fecha','dato','grupo'} → nombre de columna o None."""
    cols_norm = {c: _norm(c) for c in df.columns}
    mapping = {k: None for k in ["actor", "mesa", "nivel", "fecha", "dato", "grupo"]}

    for role, alias_list in SYNONYMS.items():
        # prioridad match exacto por normalización; luego parcial
        exact = [c for c, cn in cols_norm.items() if cn in alias_list]
        if exact:
            mapping[role] = exact[0]
            continue
        partial = [c for c, cn in cols_norm.items() if any(a in cn for a in alias_list)]
        if partial:
            mapping[role] = partial[0]

    # fallback de grupo si no existe
    if mapping["grupo"] is None:
        for c, cn in cols_norm.items():
            if any(x in cn for x in ("entidad", "secretaria", "secretaría", "dependencia", "despacho", "grupo")):
                mapping["grupo"] = c
                break
    return mapping

def auto_detect_group_fields(df: pd.DataFrame) -> List[str]:
    """Sugiere campos para agrupar (p. ej., ACTOR y/o Dependencia)."""
    mapping = guess_mapping(df)
    fields = []
    if mapping.get("actor"):
        fields.append(mapping["actor"])
    if mapping.get("grupo") and mapping["grupo"] not in fields:
        fields.append(mapping["grupo"])
    return fields[:2]

# ================= Fechas ================= #

def parse_date(val) -> Optional[pd.Timestamp]:
    if pd.isna(val):
        return None
    # Soporte números de serie Excel
    if isinstance(val, (int, float)):
        try:
            return pd.to_datetime("1899-12-30") + pd.to_timedelta(int(val), unit="D")
        except Exception:
            pass

    for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y", "%Y/%m/%d", "%m/%d/%Y", "%d.%m.%Y"):
        try:
            return pd.to_datetime(datetime.strptime(str(val).strip(), fmt))
        except Exception:
            continue
    try:
        return pd.to_datetime(val, dayfirst=True, errors="coerce")
    except Exception:
        return None

def format_date_dmy(ts: Optional[pd.Timestamp]) -> str:
    if ts is None or pd.isna(ts):
        return ""
    try:
        return pd.to_datetime(ts).strftime("%d/%m/%Y")
    except Exception:
        return ""

# ================= Preparación del DataFrame ================= #

def prepare_dataframe(df: pd.DataFrame, mapping: Dict[str, str]) -> pd.DataFrame:
    """
    Devuelve DataFrame estándar:
    ACTOR, GRUPO, MESA, NIVEL, FECHA, DATO, _FECHA_TS, FECHA_FMT
    """
    for key in ["actor", "mesa", "nivel", "fecha", "dato"]:
        if key not in mapping or mapping[key] not in df.columns:
            raise ValueError(f"Columna requerida no encontrada para '{key}'.")

    cols = [mapping["actor"], mapping.get("grupo")] if mapping.get("grupo") else [mapping["actor"]]
    cols += [mapping["mesa"], mapping["nivel"], mapping["fecha"], mapping["dato"]]
    work = df[cols].copy()
    out_cols = ["ACTOR"] + (["GRUPO"] if mapping.get("grupo") else [] ) + ["MESA", "NIVEL", "FECHA", "DATO"]
    work.columns = out_cols

    work["_FECHA_TS"] = work["FECHA"].apply(parse_date)
    work["FECHA_FMT"] = work["_FECHA_TS"].apply(format_date_dmy)
    return work

# ================= DOCX: localizar y escribir tabla ================= #

def list_candidate_tables(doc: Document) -> List[int]:
    """Devuelve índices de tablas candidatas (4 columnas)."""
    return [i for i, t in enumerate(doc.tables) if len(t.columns) == 4]

def _header_matches(table) -> bool:
    try:
        hdr_cells = table.rows[0].cells
    except Exception:
        return False
    headers = [_norm(c.text) for c in hdr_cells]
    return all(any(h == _norm(exp) or exp in h for h in headers) for exp in EXPECTED_HEADERS)

def find_target_table(doc: Document, prefer_index: int | None = None):
    """Elige la tabla cuyo encabezado se parezca; si no, última de 4 columnas; permite preferir índice."""
    if prefer_index is not None and 0 <= prefer_index < len(doc.tables):
        t = doc.tables[prefer_index]
        if len(t.columns) == 4:
            return t
    candidate = None
    for t in doc.tables:
        if _header_matches(t):
            return t
        if len(t.columns) == 4:
            candidate = t
    return candidate

def clear_table_keep_header(table) -> None:
    while len(table.rows) > 1:
        tbl = table._tbl
        tbl.remove(table.rows[-1]._tr)

def fill_table(table, rows: List[List[str]]) -> None:
    for r in rows:
        row = table.add_row()
        for i in range(min(4, len(r))):
            row.cells[i].text = "" if r[i] is None else str(r[i])

# ================= Placeholders desde Excel ================= #

def read_placeholders_from_excel(xls_bytes: bytes) -> Dict[str, Dict[str, str]]:
    """
    Lee hoja 'Placeholders' con columnas: Grupo, Llave, Valor
    Retorna: { 'NombreGrupo': {'DESTINATARIO':'...', 'CARGO':'...'} }
    """
    try:
        xls = pd.ExcelFile(io.BytesIO(xls_bytes))
    except Exception:
        return {}
    if "Placeholders" not in xls.sheet_names:
        return {}
    df = pd.read_excel(xls, "Placeholders")
    required = ["Grupo", "Llave", "Valor"]
    if not all(c in df.columns for c in required):
        return {}
    out: Dict[str, Dict[str, str]] = {}
    for _, r in df.iterrows():
        g = str(r["Grupo"]).strip()
        k = str(r["Llave"]).strip()
        v = "" if pd.isna(r["Valor"]) else str(r["Valor"])
        if not g or not k:
            continue
        out.setdefault(g, {})[k] = v
    return out
