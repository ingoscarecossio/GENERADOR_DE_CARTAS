
# -*- coding: utf-8 -*-
"""
Funcionalidades (lógica de negocio)
- Construir carta por grupo a partir de plantilla + filas
- Reemplazo de placeholders robusto (normaliza runs)
- Generación de índice y ZIP
"""
from __future__ import annotations

import io
import re
import zipfile
from typing import Dict, List, Optional, Tuple

import pandas as pd
from docx import Document

from .backend import find_target_table, clear_table_keep_header, fill_table, slugify

# --------- Reemplazo de placeholders --------- #

def _replace_placeholders_runsafe(doc: Document, mapping: Dict[str, str], normalize_runs: bool = True) -> None:
    tokens = {f"{{{{{k}}}}}": str(v) for k, v in mapping.items()}

    def _replace_in_paragraph(p):
        if not tokens:
            return
        if normalize_runs:
            full = "".join(run.text for run in p.runs)
            changed = False
            for k, v in tokens.items():
                if k in full:
                    full = full.replace(k, v)
                    changed = True
            if changed:
                for run in p.runs[1:]:
                    run.text = ""
                if p.runs:
                    p.runs[0].text = full
                else:
                    p.add_run(full)
        else:
            for run in p.runs:
                txt = run.text
                for k, v in tokens.items():
                    if k in txt:
                        txt = txt.replace(k, v)
                run.text = txt

    for p in doc.paragraphs:
        _replace_in_paragraph(p)

    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph(p)

# --------- Constructor de carta --------- #

def build_letter_bytes(template_bytes: bytes, rows: List[List[str]], placeholders: Optional[Dict[str, str]] = None, table_index: int | None = None) -> bytes:
    doc = Document(io.BytesIO(template_bytes))
    table = find_target_table(doc, prefer_index=table_index)
    if table is None:
        raise RuntimeError("No se encontró una tabla válida (4 columnas) en la plantilla.")
    clear_table_keep_header(table)
    fill_table(table, rows)
    if placeholders:
        _replace_placeholders_runsafe(doc, placeholders, normalize_runs=True)
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()

def _rows_from_group(gdf: pd.DataFrame) -> List[List[str]]:
    return [
        [
            "" if pd.isna(r.MESA) else str(r.MESA),
            "" if pd.isna(r.NIVEL) else str(r.NIVEL),
            "" if pd.isna(r.FECHA_FMT) else str(r.FECHA_FMT),
            "" if pd.isna(r.DATO) else str(r.DATO),
        ]
        for r in gdf.itertuples(index=False)
    ]

# --------- Generación por grupo --------- #

def generate_letters_per_group(
    work_df: pd.DataFrame,
    template_bytes: bytes,
    group_field: str = "ACTOR",  # por defecto ACTOR
    placeholders_per_group: Optional[Dict[str, Dict[str, str]]] = None,
    table_index: int | None = None,
    newest_first: bool = True,
) -> Tuple[Dict[str, bytes], Dict[str, str], pd.DataFrame]:
    """
    Devuelve (outputs, errors, index_df)
      outputs = {filename: bytes}
      index_df = resumen con conteo por grupo
    """
    outputs: Dict[str, bytes] = {}
    errors: Dict[str, str] = {}
    summary_rows: List[List[str]] = []

    # Orden global por fecha según preferencia
    work_df = work_df.sort_values("_FECHA_TS", ascending=not newest_first, na_position="last")

    for grp, gdf in work_df.groupby(group_field, dropna=False):
        grp_name = "(Sin grupo)" if pd.isna(grp) else str(grp)
        filas = _rows_from_group(gdf)
        try:
            placeholders = None
            if placeholders_per_group and grp_name in placeholders_per_group:
                placeholders = placeholders_per_group[grp_name]
            doc_bytes = build_letter_bytes(template_bytes, filas, placeholders, table_index=table_index)
            fname = f"CARTA_{slugify(grp_name)}.docx"
            outputs[fname] = doc_bytes
            summary_rows.append([grp_name, len(filas)])
        except Exception as e:
            errors[grp_name] = str(e)

    index_df = pd.DataFrame(summary_rows, columns=["Grupo", "Registros"]).sort_values("Grupo").reset_index(drop=True)
    return outputs, errors, index_df

# --------- Índice (Excel) --------- #

def build_index_sheet(index_df: pd.DataFrame, errors: Dict[str, str]) -> bytes:
    out = io.BytesIO()
    with pd.ExcelWriter(out, engine="openpyxl") as xlw:
        index_df.to_excel(xlw, sheet_name="Resumen", index=False)
        if errors:
            pd.DataFrame([{"Grupo": g, "Error": e} for g, e in errors.items()]).to_excel(xlw, sheet_name="Errores", index=False)
    return out.getvalue()

# --------- ZIP --------- #

def make_zip(outputs: Dict[str, bytes]) -> bytes:
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
        for fname, data in outputs.items():
            zf.writestr(fname, data)
    return buf.getvalue()
