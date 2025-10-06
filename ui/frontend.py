
# -*- coding: utf-8 -*-
"""
Frontend (Streamlit)
- Subida de plantilla y Excel
- Mapeo asistido con columna ACTOR explícita
- Filtro por ACTOR (multi-selección)
- Selección de tabla en la plantilla (si hay varias)
- Ordenamiento: "más reciente primero" por defecto
- Generación por grupo (default: ACTOR) + ZIP + índice
"""
from __future__ import annotations

from typing import Dict, Optional, List

import pandas as pd
import streamlit as st
from docx import Document

from core.backend import (
    guess_mapping, prepare_dataframe, auto_detect_group_fields,
    read_placeholders_from_excel, list_candidate_tables
)
from core.funcionalidades import generate_letters_per_group, make_zip, build_index_sheet


def run_app() -> None:
    st.set_page_config(page_title="Cartas por ACTOR — Mesas POT", layout="wide")
    st.title("Generador de Cartas filtradas por ACTOR")
    st.caption("Carga tu Excel base, filtra por ACTOR y genera cartas con la tabla: Nombre de la mesa | Nivel | Fecha | Dato transformador.")

    with st.sidebar:
        st.header("Archivos de entrada")
        tpl_file = st.file_uploader("Plantilla de carta (.docx)", type=["docx"], accept_multiple_files=False)
        xls_file = st.file_uploader("Base de datos (.xlsx/.xls)", type=["xlsx", "xls"], accept_multiple_files=False)
        st.markdown("---")
        st.header("Opciones")
        newest_first = st.checkbox("Ordenar por fecha: MÁS RECIENTE primero", value=True)
        drop_duplicates = st.checkbox("Eliminar registros duplicados exactos", value=True)
        remove_empty_mesa = st.checkbox("Descartar filas sin 'Nombre de la mesa'", value=True)

    if not (tpl_file and xls_file):
        st.info("Sube la plantilla y el Excel para comenzar.")
        return

    # Leer Excel
    try:
        df = pd.read_excel(xls_file)
    except Exception as e:
        st.error(f"No se pudo leer el Excel: {e}")
        return

    st.subheader("Mapeo de columnas")
    auto_map = guess_mapping(df)

    c1, c2, c3 = st.columns(3)
    with c1:
        actor_col = st.selectbox("Columna: ACTOR (filtro y agrupación)",
                                 options=[None] + list(df.columns),
                                 index=(list(df.columns).index(auto_map.get("actor")) + 1) if auto_map.get("actor") in df.columns else 0)
        mesa_col = st.selectbox("Columna: Nombre de la mesa",
                                 options=[None] + list(df.columns),
                                 index=(list(df.columns).index(auto_map.get("mesa")) + 1) if auto_map.get("mesa") in df.columns else 0)
    with c2:
        nivel_col = st.selectbox("Columna: Nivel",
                                 options=[None] + list(df.columns),
                                 index=(list(df.columns).index(auto_map.get("nivel")) + 1) if auto_map.get("nivel") in df.columns else 0)
        fecha_col = st.selectbox("Columna: Fecha",
                                 options=[None] + list(df.columns),
                                 index=(list(df.columns).index(auto_map.get("fecha")) + 1) if auto_map.get("fecha") in df.columns else 0)
    with c3:
        dato_col  = st.selectbox("Columna: Dato transformador",
                                 options=[None] + list(df.columns),
                                 index=(list(df.columns).index(auto_map.get("dato")) + 1) if auto_map.get("dato") in df.columns else 0)
        group_field = st.selectbox("Campo para agrupar cartas", options=[actor_col] + [c for c in df.columns if c != actor_col])

    required = {"actor": actor_col, "mesa": mesa_col, "nivel": nivel_col, "fecha": fecha_col, "dato": dato_col}
    missing = [k for k, v in required.items() if not v]
    if missing:
        st.warning("Selecciona todas las columnas requeridas para continuar.")
        return

    # Preparación estándar
    mapping = {
        "actor": actor_col,
        "mesa": mesa_col,
        "nivel": nivel_col,
        "fecha": fecha_col,
        "dato": dato_col,
        # 'grupo' es opcional; si eliges otro agrupador diferente a ACTOR, lo fijamos más abajo
    }
    if group_field and group_field != actor_col:
        mapping["grupo"] = group_field

    work = prepare_dataframe(df, mapping)

    # Filtro por ACTOR (multi)
    st.subheader("Filtro por ACTOR")
    unique_actors = sorted([a for a in work["ACTOR"].dropna().astype(str).unique() if a.strip()])
    selected_actors = st.multiselect("Selecciona uno o varios ACTORES", options=unique_actors, default=unique_actors)
    if selected_actors:
        work = work[work["ACTOR"].astype(str).isin(selected_actors)]

    # Limpiezas/orden
    if remove_empty_mesa:
        work = work[work["MESA"].notna() & (work["MESA"].astype(str).str.strip() != "")]
    if drop_duplicates:
        work = work.drop_duplicates(subset=["ACTOR", "MESA", "NIVEL", "FECHA", "DATO"])

    # Vista previa (siempre ordenar por preferencia, más reciente primero por defecto)
    work = work.sort_values("_FECHA_TS", ascending=not newest_first, na_position="last")

    with st.expander("Vista previa (100 filas, ya filtradas y ordenadas)"):
        st.dataframe(work.drop(columns=["_FECHA_TS"]).head(100), use_container_width=True)

    # Selección de tabla de la plantilla (si hay varias)
    st.subheader("Selección de tabla en la plantilla")
    try:
        from docx import Document
        tpl_bytes = tpl_file.read()
        doc = Document(io.BytesIO(tpl_bytes))
        candidates = list(range(len(doc.tables)))
        four_col_candidates = [i for i, t in enumerate(doc.tables) if len(t.columns) == 4]
        table_index = st.selectbox("Elige la tabla de 4 columnas donde se llenarán los datos",
                                   options=four_col_candidates if four_col_candidates else candidates,
                                   index=0)
        # guardar bytes otra vez para el uso posterior
        tpl_bytes_final = io.BytesIO()
        doc.save(tpl_bytes_final)
        template_bytes = tpl_bytes_final.getvalue()
    except Exception:
        st.warning("No se pudo inspeccionar la plantilla; se usará la detección automática de tabla.")
        template_bytes = tpl_file.getvalue() if hasattr(tpl_file, "getvalue") else tpl_file.read()
        table_index = None

    # Placeholders (opcional desde hoja)
    try:
        xls_bytes = xls_file.getvalue() if hasattr(xls_file, "getvalue") else xls_file.read()
        placeholders_per_group = read_placeholders_from_excel(xls_bytes)
        if placeholders_per_group:
            st.caption(f"Placeholders detectados para {len(placeholders_per_group)} grupo(s).")
    except Exception:
        placeholders_per_group = None

    st.subheader("Generación de cartas")
    if st.button("Generar DOCX por grupo"):
        outputs, errors, index_df = generate_letters_per_group(
            work_df=work,
            template_bytes=template_bytes,
            group_field=("GRUPO" if ("GRUPO" in work.columns and group_field != actor_col) else "ACTOR"),
            placeholders_per_group=placeholders_per_group,
            table_index=table_index,
            newest_first=newest_first
        )

        st.success(f"Cartas generadas: {len(outputs)}")
        if not index_df.empty:
            st.caption("Resumen por grupo")
            st.dataframe(index_df, use_container_width=True)

        if errors:
            st.warning("Se presentaron errores en algunos grupos:")
            for g, e in errors.items():
                st.write(f"- **{g}**: {e}")

        # Descargas
        for fname, data in outputs.items():
            st.download_button(
                label=f"Descargar {fname}",
                data=data,
                file_name=fname,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

        if outputs:
            from core.funcionalidades import make_zip, build_index_sheet
            zip_bytes = make_zip(outputs)
            st.download_button(
                label="Descargar todas las cartas (.zip)",
                data=zip_bytes,
                file_name="cartas_por_grupo.zip",
                mime="application/zip",
            )
            idx_bytes = build_index_sheet(index_df, errors)
            st.download_button(
                label="Descargar índice (Excel)",
                data=idx_bytes,
                file_name="indice_cartas.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
